import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from jinja2 import Template

# Función para cargar la plantilla
def cargar_plantilla(ruta):
    with open(ruta, 'r', encoding='utf-8') as file:
        return file.read()

# Función para dar formato a los encabezados
def agregar_encabezado(doc, texto):
    encabezado = doc.add_paragraph()
    run = encabezado.add_run(texto)
    run.bold = True
    run.font.size = Pt(14)
    encabezado.alignment = WD_ALIGN_PARAGRAPH.LEFT

# Función para agregar texto normal con formato específico
def agregar_texto(doc, texto):
    parrafo = doc.add_paragraph(texto)
    run = parrafo.runs[0]
    run.font.size = Pt(11)

# Función para generar el archivo .docx
def generar_docx(plantilla, datos):
    # Crear un objeto de Jinja2 Template
    template = Template(plantilla)
    
    # Rellenar la plantilla con los datos
    contenido = template.render(datos)
    
    # Crear el documento docx
    doc = Document()

    # Cargar el contenido procesado por Jinja2
    lineas = contenido.split('\n')

    # Formato específico para los encabezados y párrafos
    for linea in lineas:
        if "Introducción" in linea or "Conclusiones" in linea or "Materiales y métodos" in linea or "Resultados y discusión" in linea:
            agregar_encabezado(doc, linea.strip())  # Encabezados en negrita, 14 puntos
        else:
            agregar_texto(doc, linea.strip())  # Texto normal en 11 puntos
    
    # Guardar el documento
    archivo_docx = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Documentos de Word", "*.docx")])
    if archivo_docx:
        doc.save(archivo_docx)
        messagebox.showinfo("Éxito", "El archivo ha sido guardado correctamente.")

# Validar campos
def validar_datos(datos):
    for key, value in datos.items():
        if not value:
            messagebox.showerror("Error", f"El campo '{key}' no puede estar vacío.")
            return False
    return True



# Función para obtener los datos de la interfaz y generar el documento
def crear_documento():
    datos = {
        "titulo_espanol": entrada_titulo_espanol.get(),
        "titulo_ingles": entrada_titulo_ingles.get(),
        "autores": entrada_autores.get(),
        "afiliacion_1": entrada_afiliacion.get(),
        "resumen": entrada_resumen.get(),
        "palabras_clave": entrada_palabras.get(),
        "abstract": entrada_abstract.get(),
        "keywords": entrada_keywords.get(),
        "introduccion": entrada_introduccion.get(),
        "materiales_metodos": entrada_materiales_metodos.get(),
        "resultados_discusion": entrada_resultados_discusion.get(),
        "conclusiones": entrada_conclusiones.get()
    }
    
    # Cargar la plantilla
    ruta_plantilla = filedialog.askopenfilename(filetypes=[("Archivos de texto", "*.txt")])
    if not ruta_plantilla:
        return
    
    plantilla = cargar_plantilla(ruta_plantilla)
    
    # Generar el documento
    generar_docx(plantilla, datos)

# Configuración de la ventana de la interfaz gráfica
ventana = tk.Tk()
ventana.title("Generador de Documentos .docx")
ventana.geometry("600x800")

# Etiquetas y campos de entrada
tk.Label(ventana, text="Título en Español").pack()
entrada_titulo_espanol = tk.Entry(ventana, width=60)
entrada_titulo_espanol.pack()

tk.Label(ventana, text="Título en Inglés").pack()
entrada_titulo_ingles = tk.Entry(ventana, width=60)
entrada_titulo_ingles.pack()

tk.Label(ventana, text="Autores").pack()
entrada_autores = tk.Entry(ventana, width=60)
entrada_autores.pack()

tk.Label(ventana, text="Afiliación Institucional").pack()
entrada_afiliacion = tk.Entry(ventana, width=60)
entrada_afiliacion.pack()

tk.Label(ventana, text="Resumen").pack()
entrada_resumen = tk.Entry(ventana, width=60)
entrada_resumen.pack()

tk.Label(ventana, text="Palabras Clave").pack()
entrada_palabras = tk.Entry(ventana, width=60)
entrada_palabras.pack()

tk.Label(ventana, text="Abstract (Inglés)").pack()
entrada_abstract = tk.Entry(ventana, width=60)
entrada_abstract.pack()

tk.Label(ventana, text="Keywords (Inglés)").pack()
entrada_keywords = tk.Entry(ventana, width=60)
entrada_keywords.pack()

tk.Label(ventana, text="Introducción").pack()
entrada_introduccion = tk.Entry(ventana, width=60)
entrada_introduccion.pack()

tk.Label(ventana, text="Materiales y Métodos").pack()
entrada_materiales_metodos = tk.Entry(ventana, width=60)
entrada_materiales_metodos.pack()

tk.Label(ventana, text="Resultados y Discusión").pack()
entrada_resultados_discusion = tk.Entry(ventana, width=60)
entrada_resultados_discusion.pack()

tk.Label(ventana, text="Conclusiones").pack()
entrada_conclusiones = tk.Entry(ventana, width=60)
entrada_conclusiones.pack()

# Botón para generar el documento
tk.Button(ventana, text="Generar Documento", command=crear_documento).pack(pady=20)

# Iniciar la aplicación
ventana.mainloop()
