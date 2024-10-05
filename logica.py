# logica.py
import json
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from jinja2 import Template
from tkinter import messagebox
from reportlab.pdfgen import canvas

# Función para cargar la plantilla desde un archivo
def cargar_plantilla(ruta):
    with open(ruta, 'r', encoding='utf-8') as file:
        return file.read()

# Función para validar los datos ingresados por el usuario
def validar_datos(datos):
    for key, value in datos.items():
        if not value:
            messagebox.showerror("Error", f"El campo '{key}' no puede estar vacío.")
            return False
    return True

# Función para dar formato a los encabezados en el documento
def agregar_encabezado(doc, texto):
    encabezado = doc.add_paragraph()
    run = encabezado.add_run(texto)
    run.bold = True
    run.font.size = Pt(14)
    encabezado.alignment = WD_ALIGN_PARAGRAPH.LEFT

# Función para agregar texto normal con formato específico
def agregar_texto(doc, texto):
    parrafo = doc.add_paragraph(texto)
    run = parrafo.runs[0]
    run.font.size = Pt(11)

# Función para generar el archivo .docx
def generar_docx(plantilla, datos, archivo_docx):
    template = Template(plantilla)
    contenido = template.render(datos)
    doc = Document()

    # Dividir el contenido en líneas y aplicar el formato adecuado
    lineas = contenido.split('\n')
    for linea in lineas:
        if "Introducción" in linea or "Conclusiones" in linea or "Materiales y métodos" in linea or "Resultados y discusión" in linea:
            agregar_encabezado(doc, linea.strip())
        else:
            agregar_texto(doc, linea.strip())

    # Guardar el documento en la ruta especificada
    if archivo_docx:
        doc.save(archivo_docx)
        messagebox.showinfo("Éxito", "El archivo ha sido guardado correctamente.")

# Función para convertir un archivo .docx a .pdf
def convertir_a_pdf(archivo_docx, archivo_pdf):
    if not os.path.exists(archivo_docx):
        messagebox.showerror("Error", "El archivo .docx no existe.")
        return

    # Crear un canvas para el PDF
    pdf = canvas.Canvas(archivo_pdf)

    # Leer el documento .docx y escribir su contenido en el PDF
    doc = Document(archivo_docx)
    y = 800  # Posición Y inicial

    for para in doc.paragraphs:
        if y < 50:
            pdf.showPage()
            y = 800
        pdf.drawString(50, y, para.text)
        y -= 20  # Espacio entre párrafos

    # Guardar el PDF
    pdf.save()
    messagebox.showinfo("Éxito", "El archivo PDF ha sido guardado correctamente.")
